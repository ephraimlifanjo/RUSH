import base64,html as html_lib,io,re
from pathlib import Path
import engine
from bs4 import BeautifulSoup,NavigableString,Tag
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches
from PIL import Image
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate,Paragraph,Spacer,Image as RLImage,Table,TableStyle

def _data_image(src):
 if not isinstance(src,str) or not src.startswith('data:image/') or ',' not in src:return None
 head,data=src.split(',',1)
 try:return base64.b64decode(data)
 except Exception:return None

def _width_percent(node,default=70):
 raw=node.get('data-rush-width') or ''
 if not raw:
  m=re.search(r'width\s*:\s*([0-9.]+)%',node.get('style',''),re.I);raw=m.group(1) if m else ''
 try:return max(10,min(100,float(raw or default)))
 except Exception:return default

def _align(para,node):
 style=(node.get('style','') if isinstance(node,Tag) else '').lower()
 if 'text-align:center' in style:para.alignment=WD_ALIGN_PARAGRAPH.CENTER
 elif 'text-align:right' in style:para.alignment=WD_ALIGN_PARAGRAPH.RIGHT
 elif 'text-align:justify' in style:para.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY

def _append_inline(para,node,bold=False,italic=False,underline=False):
 if isinstance(node,NavigableString):
  if str(node):
   r=para.add_run(str(node));r.bold=bold;r.italic=italic;r.underline=underline
  return
 if not isinstance(node,Tag):return
 b=bold or node.name in ('b','strong');i=italic or node.name in ('i','em');u=underline or node.name=='u'
 if node.name=='br':para.add_run().add_break();return
 if node.name=='img':
  blob=_data_image(node.get('src'))
  if blob:
   try:para.add_run().add_picture(io.BytesIO(blob),width=Inches(6.2*_width_percent(node)/100))
   except Exception:pass
  return
 for child in node.children:_append_inline(para,child,b,i,u)

def _blocks(soup):
 root=soup.body or soup
 for child in root.children:
  if isinstance(child,Tag):yield child

def save_docx_rich(source,dest):
 soup=BeautifulSoup(source,'html.parser');doc=Document()
 def block(node):
  if node.name in ('div','section','article'):
   for c in node.children:
    if isinstance(c,Tag):block(c)
   return
  if node.name in ('ul','ol'):
   for li in node.find_all('li',recursive=False):
    p=doc.add_paragraph(style='List Number' if node.name=='ol' else 'List Bullet');_append_inline(p,li);_align(p,li)
   return
  if node.name=='table':
   rows=node.find_all('tr');cols=max([len(r.find_all(['td','th'],recursive=False)) for r in rows]or[1]);t=doc.add_table(rows=max(1,len(rows)),cols=max(1,cols));t.style='Table Grid'
   for ri,row in enumerate(rows):
    for ci,cell in enumerate(row.find_all(['td','th'],recursive=False)):
     p=t.cell(ri,ci).paragraphs[0];p.clear();_append_inline(p,cell)
   return
  if node.name=='img':
   p=doc.add_paragraph();_append_inline(p,node);return
  style={'h1':'Title','h2':'Heading 1','h3':'Heading 2'}.get(node.name)
  p=doc.add_paragraph(style=style) if style else doc.add_paragraph();_append_inline(p,node);_align(p,node)
 for n in _blocks(soup):block(n)
 doc.core_properties.author='Nova Studio Plateformes';Path(dest).parent.mkdir(parents=True,exist_ok=True);doc.save(dest)

def _run_html(run,doc):
 text=html_lib.escape(run.text or '')
 if run.underline:text=f'<u>{text}</u>'
 if run.italic:text=f'<i>{text}</i>'
 if run.bold:text=f'<b>{text}</b>'
 out=[text] if text else []
 try:
  for blip in run._element.xpath('.//a:blip'):
   rid=blip.get(qn('r:embed'));part=doc.part.related_parts.get(rid)
   if not part:continue
   mime=getattr(part,'content_type','image/png');data=base64.b64encode(part.blob).decode('ascii');out.append(f'<img src="data:{mime};base64,{data}" data-rush-width="70" style="width:70%;height:auto">')
 except Exception:pass
 return ''.join(out)

def open_docx_rich(src):
 doc=Document(src);parts=[]
 for p in doc.paragraphs:
  inner=''.join(_run_html(r,doc) for r in p.runs) or html_lib.escape(p.text or '')
  style=(p.style.name if p.style else '').lower();tag='h1' if 'title' in style else 'h2' if 'heading 1' in style else 'h3' if 'heading 2' in style else 'p'
  align='center' if p.alignment==WD_ALIGN_PARAGRAPH.CENTER else 'right' if p.alignment==WD_ALIGN_PARAGRAPH.RIGHT else 'justify' if p.alignment==WD_ALIGN_PARAGRAPH.JUSTIFY else 'left';parts.append(f'<{tag} style="text-align:{align}">{inner}</{tag}>')
 for table in doc.tables:
  rows=[]
  for row in table.rows:rows.append('<tr>'+''.join(f'<td>{html_lib.escape(cell.text)}</td>' for cell in row.cells)+'</tr>')
  parts.append('<table>'+''.join(rows)+'</table>')
 engine.out(html=''.join(parts),format='.docx')

def _paragraph_html(node):return ''.join(str(x) for x in node.contents)
def save_pdf_rich(source,dest):
 soup=BeautifulSoup(source,'html.parser');styles=getSampleStyleSheet();story=[]
 for n in _blocks(soup):
  if n.name=='img':
   blob=_data_image(n.get('src'))
   if blob:
    im=Image.open(io.BytesIO(blob));ratio=im.height/max(1,im.width);w=6.5*inch*_width_percent(n)/100;story.append(RLImage(io.BytesIO(blob),width=w,height=w*ratio));story.append(Spacer(1,8))
  elif n.name=='table':
   data=[[c.get_text(' ',strip=True) for c in r.find_all(['td','th'],recursive=False)] for r in n.find_all('tr')]
   if data:
    cols=max(len(r) for r in data);data=[r+['']*(cols-len(r)) for r in data];t=Table(data,repeatRows=1);t.setStyle(TableStyle([('GRID',(0,0),(-1,-1),.5,colors.grey),('PADDING',(0,0),(-1,-1),5)]));story.extend([t,Spacer(1,8)])
  elif n.name in ('ul','ol'):
   for li in n.find_all('li',recursive=False):story.append(Paragraph(('• ' if n.name=='ul' else '')+html_lib.escape(li.get_text(' ',strip=True)),styles['BodyText']))
  else:
   sty=styles['Title'] if n.name=='h1' else styles['Heading1'] if n.name=='h2' else styles['Heading2'] if n.name=='h3' else styles['BodyText'];story.append(Paragraph(_paragraph_html(n) or '&nbsp;',sty));story.append(Spacer(1,4))
 Path(dest).parent.mkdir(parents=True,exist_ok=True);SimpleDocTemplate(dest,pagesize=A4,rightMargin=54,leftMargin=54,topMargin=54,bottomMargin=54).build(story or [Paragraph(' ',styles['BodyText'])])

def op_doc_open(p):
 src=engine.req(p);e=Path(src).suffix.lower()
 if e=='.docx':return open_docx_rich(src)
 return engine.op_doc_open(p)

def op_doc_save(p):
 dest=str(p.get('output')or'');source=str(p.get('html')or'');fmt=(p.get('format')or Path(dest).suffix or'.docx').lower()
 if not dest:raise ValueError('Output path required')
 if fmt=='.docx':save_docx_rich(source,dest);engine.out(output=dest,format=fmt);return
 if fmt=='.pdf':save_pdf_rich(source,dest);engine.out(output=dest,format=fmt);return
 return engine.op_doc_save(p)
