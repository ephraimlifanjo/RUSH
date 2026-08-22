import os, json, html as html_lib
from pathlib import Path

try:
    from bs4 import BeautifulSoup, NavigableString
except Exception:
    BeautifulSoup = NavigableString = None

try:
    from docx import Document
except Exception:
    Document = None

try:
    from pyhanko.sign import signers
    from pyhanko.pdf_utils.incremental_writer import IncrementalPdfFileWriter
except Exception:
    signers = IncrementalPdfFileWriter = None

try:
    import argostranslate.translate as argos_translate
except Exception:
    argos_translate = None


def signing_available():
    return signers is not None and IncrementalPdfFileWriter is not None


def translation_available():
    return argos_translate is not None


def op_sign_pdf(p, out):
    if not signing_available():
        raise ValueError('pyHanko is unavailable. Install the RUSH signing component first.')
    src = p.get('input')
    dest = p.get('output')
    cert = p.get('certificate')
    password = str(p.get('password') or '')
    field_name = str(p.get('fieldName') or 'RUSHSignature1')
    reason = str(p.get('reason') or 'Document approved in RUSH Office Suite')
    location = str(p.get('location') or '')
    if not src or not Path(src).exists():
        raise ValueError('Input PDF not found')
    if not dest:
        raise ValueError('Output path required')
    if not cert or not Path(cert).exists() or Path(cert).suffix.lower() not in ('.p12','.pfx'):
        raise ValueError('Choose a PKCS#12 certificate (.p12 or .pfx)')
    signer = signers.SimpleSigner.load_pkcs12(cert, passphrase=password.encode('utf-8') if password else None)
    if signer is None:
        raise ValueError('Could not unlock the signing certificate')
    meta = signers.PdfSignatureMetadata(field_name=field_name, reason=reason, location=location or None)
    pdf_signer = signers.PdfSigner(meta, signer=signer)
    Path(dest).parent.mkdir(parents=True, exist_ok=True)
    with open(src, 'rb') as inf, open(dest, 'wb') as outf:
        writer = IncrementalPdfFileWriter(inf)
        pdf_signer.sign_pdf(writer, output=outf)
    out(output=str(dest), signed=True, standard='PDF digital signature / PAdES-capable engine')


def _translation(from_code, to_code):
    if not translation_available():
        raise ValueError('Offline translation packs are not installed. Install python/requirements-translation.txt and the required local language packs.')
    langs = argos_translate.get_installed_languages()
    source = next((x for x in langs if x.code == from_code), None)
    target = next((x for x in langs if x.code == to_code), None)
    if source is None or target is None:
        raise ValueError(f'Offline translation pack missing for {from_code} -> {to_code}')
    try:
        return source.get_translation(target)
    except Exception:
        raise ValueError(f'No installed offline translation route for {from_code} -> {to_code}')


def _translate_html(html, trans):
    if BeautifulSoup is None:
        raise ValueError('BeautifulSoup is unavailable')
    soup = BeautifulSoup(html, 'html.parser')
    for node in list(soup.find_all(string=True)):
        if not isinstance(node, NavigableString):
            continue
        parent = getattr(node, 'parent', None)
        if parent and parent.name in ('script','style','code','pre'):
            continue
        raw = str(node)
        if not raw.strip():
            continue
        lead = raw[:len(raw)-len(raw.lstrip())]
        tail = raw[len(raw.rstrip()):]
        translated = trans.translate(raw.strip())
        node.replace_with(lead + translated + tail)
    return str(soup)


def op_translate_document(p, out):
    src = p.get('input')
    dest = p.get('output')
    from_code = str(p.get('from') or 'en')
    to_code = str(p.get('to') or 'fr')
    if not src or not Path(src).exists():
        raise ValueError('Input document not found')
    if not dest:
        raise ValueError('Output path required')
    trans = _translation(from_code, to_code)
    ext = Path(src).suffix.lower()
    Path(dest).parent.mkdir(parents=True, exist_ok=True)

    if ext == '.docx':
        if Document is None:
            raise ValueError('python-docx is unavailable')
        doc = Document(src)
        for para in doc.paragraphs:
            for run in para.runs:
                if run.text.strip():
                    run.text = trans.translate(run.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if run.text.strip():
                                run.text = trans.translate(run.text)
        doc.save(dest)
        out(output=str(dest), translated=True, preserved='paragraphs, runs, tables and most DOCX styling')
        return

    if ext in ('.txt', '.rtf'):
        text = Path(src).read_text(encoding='utf-8', errors='replace')
        Path(dest).write_text(trans.translate(text), encoding='utf-8')
        out(output=str(dest), translated=True, preserved='plain text')
        return

    if ext in ('.html', '.htm'):
        html = Path(src).read_text(encoding='utf-8', errors='replace')
        Path(dest).write_text(_translate_html(html, trans), encoding='utf-8')
        out(output=str(dest), translated=True, preserved='HTML structure')
        return

    raise ValueError('Structure-preserving translation currently supports DOCX, TXT, RTF and HTML. PDF translation is intentionally not advertised as exact-layout translation yet.')


def capabilities():
    return {
        'digitalSignature': signing_available(),
        'offlineTranslation': translation_available()
    }
