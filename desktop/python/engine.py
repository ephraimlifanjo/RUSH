import sys,json,os,io,tempfile,sqlite3,re,subprocess,shutil,html as html_lib
from pathlib import Path
from pypdf import PdfReader,PdfWriter
from pypdf.generic import NameObject,DictionaryObject,ArrayObject,FloatObject
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
from PIL import Image,ImageDraw,ImageOps
import pypdfium2 as pdfium
try:
 from bs4 import BeautifulSoup
except Exception: BeautifulSoup=None
try:
 from docx import Document
 from docx.enum.text import WD_ALIGN_PARAGRAPH
except Exception: Document=WD_ALIGN_PARAGRAPH=None
try:
 from odf.opendocument import load as odt_load,OpenDocumentText
 from odf.text import P as OdtP,H as OdtH
except Exception: odt_load=OpenDocumentText=OdtP=OdtH=None
try:
 import pytesseract
except Exception: pytesseract=None
try:
 import duckdb
except Exception: duckdb=None

def payload():
 raw=sys.stdin.read().strip();return json.loads(raw) if raw else {}
def out(**kw): print(json.dumps({'ok':True,**kw},ensure_ascii=False))
def fail(e): print(json.dumps({'ok':False,'error':str(e)},ensure_ascii=False));sys.exit(2)
def req(p,k='input'):
 v=p.get(k)
 if not v or not Path(v).exists(): raise ValueError(f'Missing or inaccessible {k}')
 return str(Path(v))
def save(w,dest):
 Path(dest).parent.mkdir(parents=True,exist_ok=True)
 with open(dest,'wb') as f:w.write(f)
 return str(dest)
def copy_meta(r,w):
 try:
  if r.metadata:w.add_metadata({str(k):str(v) for k,v in r.metadata.items() if v is not None})
 except Exception:pass
def pages(spec,total):
 if spec is None or str(spec).strip().lower() in ('','all'):return list(range(total))
 if isinstance(spec,list):raw=spec
 else:raw=str(spec).replace(' ','').split(',')
 ans=[]
 for part in raw:
  if isinstance(part,int):ans.append(part-1);continue
  if not part:continue
  if '-' in str(part):
   a,b=map(int,str(part).split('-',1));a,b=min(a,b),max(a,b);ans.extend(range(a-1,b))
  else:ans.append(int(part)-1)
 seen=set();return[i for i in ans if 0<=i<total and not(i in seen or seen.add(i))]
def page_size(pg):return float(pg.mediabox.width),float(pg.mediabox.height)
def overlay(width,height,draw):
 b=io.BytesIO();c=canvas.Canvas(b,pagesize=(width,height));draw(c);c.save();b.seek(0);return PdfReader(b).pages[0]
def render(path,index,dpi=150):
 doc=pdfium.PdfDocument(path)
 try:return doc[index].render(scale=max(.5,dpi/72)).to_pil().convert('RGB')
 finally:
  try:doc.close()
  except Exception:pass

def op_info(p):
 r=PdfReader(req(p));out(pages=len(r.pages),encrypted=bool(r.is_encrypted),metadata={str(k):str(v) for k,v in(r.metadata or {}).items()},fileSize=os.path.getsize(p['input']))
def op_merge(p):
 ins=[x for x in p.get('inputs',[]) if Path(x).exists()]
 if len(ins)<2:raise ValueError('Choose at least two PDF files')
 w=PdfWriter()
 for f in ins:
  for pg in PdfReader(f).pages:w.add_page(pg)
 out(output=save(w,p['output']),pages=len(w.pages))
def op_extract(p):
 r=PdfReader(req(p));sel=pages(p.get('pages'),len(r.pages));w=PdfWriter()
 if not sel:raise ValueError('No valid pages selected')
 for i in sel:w.add_page(r.pages[i])
 copy_meta(r,w);out(output=save(w,p['output']),pages=len(w.pages))
def op_delete(p):
 r=PdfReader(req(p));rm=set(pages(p.get('pages'),len(r.pages)))
 if len(rm)>=len(r.pages):raise ValueError('Cannot delete every page')
 w=PdfWriter()
 for i,pg in enumerate(r.pages):
  if i not in rm:w.add_page(pg)
 copy_meta(r,w);out(output=save(w,p['output']),pages=len(w.pages))
def op_rotate(p):
 r=PdfReader(req(p));target=set(pages(p.get('pages','all'),len(r.pages)));deg=int(p.get('degrees',90))
 if deg not in(90,180,270,-90,-180,-270):raise ValueError('Rotation must be 90, 180 or 270')
 w=PdfWriter()
 for i,pg in enumerate(r.pages):
  if i in target:pg.rotate(deg)
  w.add_page(pg)
 copy_meta(r,w);out(output=save(w,p['output']))
def op_reorder(p):
 r=PdfReader(req(p));order=pages(p.get('order'),len(r.pages))
 if len(order)!=len(r.pages):raise ValueError('Page order must include every page once')
 w=PdfWriter()
 for i in order:w.add_page(r.pages[i])
 copy_meta(r,w);out(output=save(w,p['output']))
def op_duplicate(p):
 r=PdfReader(req(p));target=set(pages(p.get('pages'),len(r.pages)));copies=max(1,min(20,int(p.get('copies',1))));w=PdfWriter()
 for i,pg in enumerate(r.pages):
  w.add_page(pg)
  if i in target:
   for _ in range(copies):w.add_page(pg)
 copy_meta(r,w);out(output=save(w,p['output']),pages=len(w.pages))
def op_insert_blank(p):
 r=PdfReader(req(p));pos=max(0,min(len(r.pages),int(p.get('position',len(r.pages)+1))-1));count=max(1,min(50,int(p.get('count',1))));w=PdfWriter();ref=r.pages[min(max(pos-1,0),len(r.pages)-1)] if r.pages else None;ww,hh=page_size(ref) if ref else(612,792)
 for i,pg in enumerate(r.pages):
  if i==pos:
   for _ in range(count):w.add_blank_page(width=ww,height=hh)
  w.add_page(pg)
 if pos>=len(r.pages):
  for _ in range(count):w.add_blank_page(width=ww,height=hh)
 copy_meta(r,w);out(output=save(w,p['output']))
def op_crop(p):
 r=PdfReader(req(p));target=set(pages(p.get('pages','all'),len(r.pages)));m={k:max(0,float(p.get(k,0)or 0)) for k in('left','right','top','bottom')};w=PdfWriter()
 for i,pg in enumerate(r.pages):
  if i in target:
   x0=float(pg.mediabox.left)+m['left'];y0=float(pg.mediabox.bottom)+m['bottom'];x1=float(pg.mediabox.right)-m['right'];y1=float(pg.mediabox.top)-m['top']
   if x1<=x0 or y1<=y0:raise ValueError('Crop margins are too large')
   pg.cropbox.lower_left=(x0,y0);pg.cropbox.upper_right=(x1,y1)
  w.add_page(pg)
 copy_meta(r,w);out(output=save(w,p['output']))
def op_split(p):
 r=PdfReader(req(p));every=max(1,int(p.get('every',1)));folder=Path(p.get('output_dir')or Path(p['input']).parent/(Path(p['input']).stem+'_split'));folder.mkdir(parents=True,exist_ok=True);outputs=[]
 for start in range(0,len(r.pages),every):
  w=PdfWriter()
  for i in range(start,min(start+every,len(r.pages))):w.add_page(r.pages[i])
  d=folder/f'{Path(p["input"]).stem}_{start+1}-{min(start+every,len(r.pages))}.pdf';save(w,d);outputs.append(str(d))
 out(outputs=outputs,count=len(outputs))
def op_watermark(p):
 r=PdfReader(req(p));text=str(p.get('text')or'CONFIDENTIAL');opacity=max(.03,min(1,float(p.get('opacity',.18)or.18)));w=PdfWriter()
 for pg in r.pages:
  ww,hh=page_size(pg)
  ov=overlay(ww,hh,lambda c:(c.saveState(),c.setFillAlpha(opacity),c.setFillColorRGB(.82,.14,.14),c.setFont('Helvetica-Bold',max(18,min(64,ww/9))),c.translate(ww/2,hh/2),c.rotate(35),c.drawCentredString(0,0,text),c.restoreState()))
  pg.merge_page(ov);w.add_page(pg)
 copy_meta(r,w);out(output=save(w,p['output']))
def op_page_numbers(p):
 r=PdfReader(req(p));start=int(p.get('start',1));w=PdfWriter()
 for i,pg in enumerate(r.pages):
  ww,hh=page_size(pg);ov=overlay(ww,hh,lambda c,n=i+start:(c.setFont('Helvetica',9),c.setFillColorRGB(.25,.25,.25),c.drawCentredString(ww/2,20,str(n))));pg.merge_page(ov);w.add_page(pg)
 copy_meta(r,w);out(output=save(w,p['output']))
def op_encrypt(p):
 r=PdfReader(req(p));pwd=str(p.get('password')or'')
 if not pwd:raise ValueError('Password required')
 w=PdfWriter();[w.add_page(x) for x in r.pages];copy_meta(r,w);w.encrypt(pwd,str(p.get('ownerPassword')or pwd),algorithm='AES-256');out(output=save(w,p['output']))
def op_decrypt(p):
 r=PdfReader(req(p));pwd=str(p.get('password')or'')
 if r.is_encrypted and not r.decrypt(pwd):raise ValueError('Incorrect password')
 w=PdfWriter();[w.add_page(x) for x in r.pages];copy_meta(r,w);out(output=save(w,p['output']))
def op_metadata(p):
 r=PdfReader(req(p));w=PdfWriter();[w.add_page(x) for x in r.pages];meta={}
 for k,v in [('Title',p.get('title')),('Author',p.get('author')),('Subject',p.get('subject')),('Keywords',p.get('keywords'))]:
  if v:meta['/'+k]=str(v)
 w.add_metadata(meta);out(output=save(w,p['output']))
def op_remove_metadata(p):
 r=PdfReader(req(p));w=PdfWriter();[w.add_page(x) for x in r.pages];out(output=save(w,p['output']))
def op_text(p):
 r=PdfReader(req(p));chunks=[pg.extract_text()or'' for pg in r.pages];text='\n\n'.join(chunks)
 if p.get('output'):Path(p['output']).write_text(text,encoding='utf-8')
 out(text=text,characters=len(text),output=p.get('output'))
def op_images_to_pdf(p):
 ins=[x for x in p.get('inputs',[]) if Path(x).exists()]
 if not ins and p.get('input'):ins=[p['input']]
 if not ins:raise ValueError('Choose image files')
 imgs=[]
 for f in ins:
  im=Image.open(f).convert('RGB');imgs.append(im.copy());im.close()
 Path(p['output']).parent.mkdir(parents=True,exist_ok=True);imgs[0].save(p['output'],'PDF',save_all=True,append_images=imgs[1:],resolution=150);out(output=p['output'],pages=len(imgs))
def op_pdf_to_images(p):
 src=req(p);folder=Path(p.get('output_dir')or p.get('output')or(Path(src).parent/(Path(src).stem+'_images')));folder.mkdir(parents=True,exist_ok=True);fmt=str(p.get('format','png')).lower();dpi=int(p.get('dpi',150));doc=pdfium.PdfDocument(src);outputs=[]
 try:
  for i in range(len(doc)):
   im=doc[i].render(scale=dpi/72).to_pil().convert('RGB');dest=folder/f'{Path(src).stem}_{i+1}.{fmt}';im.save(dest,'JPEG' if fmt in('jpg','jpeg') else 'PNG',quality=88);outputs.append(str(dest))
 finally:doc.close()
 out(outputs=outputs,count=len(outputs))
def op_compress(p):
 r=PdfReader(req(p));w=PdfWriter()
 for pg in r.pages:
  try:pg.compress_content_streams()
  except Exception:pass
  w.add_page(pg)
 copy_meta(r,w);out(output=save(w,p['output']))
def op_grayscale(p):
 src=req(p);r=PdfReader(src);w=PdfWriter();dpi=int(p.get('dpi',130))
 for i,pg in enumerate(r.pages):
  ww,hh=page_size(pg);im=ImageOps.grayscale(render(src,i,dpi)).convert('RGB');w.add_page(overlay(ww,hh,lambda c,im=im:c.drawImage(ImageReader(im),0,0,width=ww,height=hh)))
 copy_meta(r,w);out(output=save(w,p['output']))
def op_flatten(p):
 r=PdfReader(req(p));w=PdfWriter()
 for pg in r.pages:
  try:
   if '/Annots' in pg:del pg['/Annots']
  except Exception:pass
  w.add_page(pg)
 copy_meta(r,w);out(output=save(w,p['output']))
def _rect(e,ww,hh):
 x=float(e.get('x',0));y=float(e.get('y',0));rw=float(e.get('w',0));rh=float(e.get('h',0));normalized=max(x,y,rw,rh)<=1.01
 if normalized:x*=ww;y*=hh;rw*=ww;rh*=hh
 return x,hh-y-rh,rw,rh
def op_apply_edits(p):
 src=req(p);edits=p.get('edits')or[];r=PdfReader(src);w=PdfWriter();by={}
 for e in edits:by.setdefault(max(0,int(e.get('page',1))-1),[]).append(e)
 for i,pg in enumerate(r.pages):
  es=by.get(i,[]);ww,hh=page_size(pg)
  if any(e.get('tool')=='redact' for e in es):
   im=render(src,i,170);sx=im.width/ww;sy=im.height/hh;d=ImageDraw.Draw(im)
   for e in es:
    x,y,rw,rh=_rect(e,ww,hh);top=hh-y-rh
    if e.get('tool')=='redact':d.rectangle([x*sx,top*sy,(x+rw)*sx,(top+rh)*sy],fill=(0,0,0))
   pg=overlay(ww,hh,lambda c,im=im:c.drawImage(ImageReader(im),0,0,width=ww,height=hh))
  def draw(c):
   for e in es:
    tool=e.get('tool');x,y,rw,rh=_rect(e,ww,hh)
    if tool=='redact':continue
    if tool in('whiteout','edit_text'):
     c.setFillColorRGB(1,1,1);c.rect(x,y,rw,rh,fill=1,stroke=0)
    if tool=='highlight':c.saveState();c.setFillAlpha(.32);c.setFillColorRGB(1,.9,.1);c.rect(x,y,rw,rh,fill=1,stroke=0);c.restoreState()
    if tool=='rectangle':c.setStrokeColorRGB(.82,.14,.14);c.setLineWidth(1.4);c.rect(x,y,rw,rh,fill=0,stroke=1)
    if tool=='underline':c.setStrokeColorRGB(.82,.14,.14);c.line(x,y+2,x+rw,y+2)
    if tool=='strikeout':c.setStrokeColorRGB(.82,.14,.14);c.line(x,y+rh/2,x+rw,y+rh/2)
    if tool in('add_text','edit_text'):
     c.setFillColorRGB(.1,.1,.1);c.setFont('Helvetica',max(7,float(e.get('size',11))));c.drawString(x,y+max(2,rh/3),str(e.get('text')or'Text'))
    if tool=='signature' and e.get('image') and Path(e['image']).exists():c.drawImage(e['image'],x,y,width=rw,height=rh,preserveAspectRatio=True,mask='auto')
   
  if es:
   ov=overlay(ww,hh,draw);pg.merge_page(ov)
  w.add_page(pg)
 copy_meta(r,w);out(output=save(w,p['output']),edits=len(edits))
def op_form_fields(p):
 r=PdfReader(req(p));fields=r.get_fields()or{};items=[]
 for k,v in fields.items():items.append({'name':str(k),'value':str(v.get('/V','')or''),'type':str(v.get('/FT','')or'')})
 out(fields=items,count=len(items))
def op_fill_form(p):
 r=PdfReader(req(p));w=PdfWriter();w.append_pages_from_reader(r);fields=p.get('fields')or{}
 try:
  for pg in w.pages:w.update_page_form_field_values(pg,fields,auto_regenerate=True)
 except Exception:pass
 out(output=save(w,p['output']),fields=len(fields))
def tess_ready():
 if pytesseract is None:return False
 custom=os.environ.get('RUSH_TESSERACT_EXE')
 if custom and Path(custom).exists():pytesseract.pytesseract.tesseract_cmd=custom;return True
 if os.name=='nt' and Path(r'C:\Program Files\Tesseract-OCR\tesseract.exe').exists():pytesseract.pytesseract.tesseract_cmd=r'C:\Program Files\Tesseract-OCR\tesseract.exe';return True
 return shutil.which('tesseract') is not None
def page_text(src,i,langs='eng+fra',force=False):
 r=PdfReader(src);text=r.pages[i].extract_text()or''
 if(force or len(text.strip())<30)and tess_ready():
  try:return pytesseract.image_to_string(render(src,i,150),lang=langs),True
  except Exception:pass
 return text,False
def op_search_pdf(p):
 src=req(p);q=str(p.get('query')or'').strip();langs=p.get('languages')or'eng+fra'
 if not q:raise ValueError('Search text is required')
 r=PdfReader(src);matches=[];used=False
 for i in range(len(r.pages)):
  text,u=page_text(src,i,langs);used=used or u
  if q.lower() in text.lower():
   idx=text.lower().find(q.lower());matches.append({'page':i+1,'text':text[max(0,idx-60):idx+len(q)+100].replace('\n',' ')})
 out(matches=matches,count=len(matches),ocrUsed=used)
def op_ocr_pdf(p):
 src=req(p)
 if not tess_ready():raise ValueError('Tesseract OCR is not installed or configured')
 langs=p.get('languages')or'eng+fra';doc=pdfium.PdfDocument(src);w=PdfWriter()
 try:
  for i in range(len(doc)):
   im=doc[i].render(scale=2.0).to_pil().convert('RGB');pdfbytes=pytesseract.image_to_pdf_or_hocr(im,extension='pdf',lang=langs);rr=PdfReader(io.BytesIO(pdfbytes));w.add_page(rr.pages[0])
 finally:doc.close()
 out(output=save(w,p['output']),pages=len(w.pages))

def _html_from_text(text):return ''.join(f'<p>{html_lib.escape(line) or "<br>"}</p>' for line in text.splitlines())
def find_converter():return shutil.which('soffice')or shutil.which('libreoffice')
def op_doc_open(p):
 src=req(p);e=Path(src).suffix.lower()
 if e=='.doc':
  conv=find_converter()
  if not conv:raise ValueError('Legacy .doc import needs Microsoft Word or LibreOffice installed locally. Save as .docx if no converter is available.')
  td=tempfile.mkdtemp(prefix='rush-doc-');subprocess.run([conv,'--headless','--convert-to','docx','--outdir',td,src],check=True,stdout=subprocess.PIPE,stderr=subprocess.PIPE);src=str(Path(td)/(Path(src).stem+'.docx'));e='.docx'
 if e=='.docx':
  if Document is None:raise ValueError('python-docx is unavailable')
  d=Document(src);parts=[]
  for para in d.paragraphs:
   runs=''.join(('<b>' if r.bold else '')+('<i>' if r.italic else '')+('<u>' if r.underline else '')+html_lib.escape(r.text)+('</u>' if r.underline else '')+('</i>' if r.italic else '')+('</b>' if r.bold else '') for r in para.runs)
   parts.append(f'<p>{runs}</p>')
  for table in d.tables:
   parts.append('<table>')
   for row in table.rows:parts.append('<tr>'+''.join(f'<td>{html_lib.escape(c.text)}</td>' for c in row.cells)+'</tr>')
   parts.append('</table>')
  out(html=''.join(parts),format='.docx');return
 if e=='.odt':
  if odt_load is None:raise ValueError('odfpy is unavailable')
  d=odt_load(src);texts=[]
  for n in d.getElementsByType(OdtH):texts.append(f'<h2>{html_lib.escape("".join(x.data for x in n.childNodes if hasattr(x,"data")))}</h2>')
  for n in d.getElementsByType(OdtP):texts.append(f'<p>{html_lib.escape("".join(x.data for x in n.childNodes if hasattr(x,"data")))}</p>')
  out(html=''.join(texts),format='.odt');return
 raw=Path(src).read_text(encoding='utf-8' if e!='.rtf' else 'latin-1',errors='replace')
 if e=='.rtf':raw=re.sub(r'\\[a-z]+-?\d* ?|[{}]','',raw)
 out(html=raw if e in('.html','.htm') else _html_from_text(raw),format=e)
def _plain(html):return BeautifulSoup(html,'html.parser').get_text('\n') if BeautifulSoup else re.sub(r'<[^>]+>','',html)
def save_docx(html,dest):
 if Document is None:raise ValueError('python-docx unavailable')
 d=Document();soup=BeautifulSoup(html,'html.parser') if BeautifulSoup else None
 if soup:
  for n in soup.find_all(['h1','h2','h3','p','li','table'],recursive=True):
   if n.name=='table':
    rows=n.find_all('tr',recursive=False);cols=max([len(x.find_all(['td','th'],recursive=False)) for x in rows]or[1]);t=d.add_table(rows=len(rows),cols=cols);t.style='Table Grid'
    for ri,row in enumerate(rows):
     for ci,c in enumerate(row.find_all(['td','th'],recursive=False)):t.cell(ri,ci).text=c.get_text(' ',strip=True)
   elif n.find_parent('table') is None:
    st={'h1':'Title','h2':'Heading 1','h3':'Heading 2','li':'List Bullet'}.get(n.name);para=d.add_paragraph(style=st) if st else d.add_paragraph();para.add_run(n.get_text('',strip=False))
 else:d.add_paragraph(_plain(html))
 d.core_properties.author='Nova Studio Plateformes';Path(dest).parent.mkdir(parents=True,exist_ok=True);d.save(dest)
def op_doc_save(p):
 dest=str(p.get('output')or'');html=str(p.get('html')or'');fmt=(p.get('format')or Path(dest).suffix or'.docx').lower()
 if not dest:raise ValueError('Output path required')
 if fmt=='.docx':save_docx(html,dest)
 elif fmt=='.odt':
  if OpenDocumentText is None:raise ValueError('odfpy unavailable')
  d=OpenDocumentText();soup=BeautifulSoup(html,'html.parser') if BeautifulSoup else None
  for line in(soup.stripped_strings if soup else _plain(html).splitlines()):d.text.addElement(OdtP(text=str(line)))
  d.save(dest)
 elif fmt=='.rtf':Path(dest).write_text('{\\rtf1\\ansi '+_plain(html).replace('\\','\\\\').replace('{','\\{').replace('}','\\}').replace('\n','\\par ')+'}',encoding='latin-1',errors='replace')
 elif fmt=='.txt':Path(dest).write_text(_plain(html),encoding='utf-8')
 elif fmt in('.html','.htm'):Path(dest).write_text(html,encoding='utf-8')
 elif fmt=='.pdf':
  text=_plain(html);c=canvas.Canvas(dest,pagesize=(595,842));y=790;c.setFont('Helvetica',11)
  for line in text.splitlines():
   if y<50:c.showPage();c.setFont('Helvetica',11);y=790
   c.drawString(50,y,line[:105]);y-=16
  c.save()
 elif fmt=='.doc':raise ValueError('Legacy .doc is import-only. Save as .docx.')
 else:raise ValueError('Unsupported document format')
 out(output=dest,format=fmt)
def index_db():return Path(os.environ.get('RUSH_USER_DATA')or tempfile.gettempdir())/'library.sqlite3'
def ensure_db():
 db=sqlite3.connect(index_db());db.execute('CREATE TABLE IF NOT EXISTS documents(path TEXT PRIMARY KEY,name TEXT,ext TEXT,size INTEGER,modified REAL,pages INTEGER,words INTEGER,ocr INTEGER DEFAULT 0)')
 try:db.execute("CREATE VIRTUAL TABLE IF NOT EXISTS documents_fts USING fts5(path UNINDEXED,content,tokenize='unicode61')")
 except sqlite3.OperationalError:db.execute('CREATE TABLE IF NOT EXISTS documents_fts(path TEXT PRIMARY KEY,content TEXT)')
 db.commit();return db
def text_document(path,ocr=False,langs='eng+fra'):
 e=Path(path).suffix.lower()
 if e=='.pdf':
  r=PdfReader(path);arr=[];used=False
  for i in range(len(r.pages)):
   t,u=page_text(path,i,langs,force=ocr);arr.append(t);used|=u
  return '\n\n'.join(arr),len(r.pages),used
 if e=='.docx' and Document:
  d=Document(path);return '\n'.join(x.text for x in d.paragraphs),1,False
 try:return Path(path).read_text(encoding='utf-8',errors='replace'),1,False
 except:return '',1,False
def op_index_paths(p):
 paths=[x for x in p.get('paths',[]) if x and Path(x).exists()][:50000];db=ensure_db();count=0;words=0;ocrn=0
 for f in paths:
  try:
   t,pg,u=text_document(f,bool(p.get('ocr')),p.get('languages')or'eng+fra');st=Path(f).stat();wc=len(re.findall(r'\S+',t));words+=wc;ocrn+=int(u)
   db.execute('INSERT INTO documents(path,name,ext,size,modified,pages,words,ocr) VALUES(?,?,?,?,?,?,?,?) ON CONFLICT(path) DO UPDATE SET name=excluded.name,ext=excluded.ext,size=excluded.size,modified=excluded.modified,pages=excluded.pages,words=excluded.words,ocr=excluded.ocr',(f,Path(f).name,Path(f).suffix.lower(),st.st_size,st.st_mtime,pg,wc,int(u)));db.execute('DELETE FROM documents_fts WHERE path=?',(f,));db.execute('INSERT INTO documents_fts(path,content) VALUES(?,?)',(f,t));count+=1
   if count%25==0:db.commit()
  except Exception:pass
 db.commit();db.close();out(indexed=count,words=words,ocrDocuments=ocrn,db=str(index_db()))
def op_search_index(p):
 if not index_db().exists():out(results=[],count=0);return
 q=str(p.get('query')or'').strip();limit=max(1,min(500,int(p.get('limit',100))));db=sqlite3.connect(index_db());db.row_factory=sqlite3.Row
 try:rows=db.execute("SELECT d.*,snippet(documents_fts,1,'<mark>','</mark>',' … ',18) snippet FROM documents_fts JOIN documents d ON d.path=documents_fts.path WHERE documents_fts MATCH ? LIMIT ?",(q,limit)).fetchall()
 except sqlite3.OperationalError:rows=db.execute('SELECT d.*,substr(f.content,1,220) snippet FROM documents_fts f JOIN documents d ON d.path=f.path WHERE f.content LIKE ? LIMIT ?',('%'+q+'%',limit)).fetchall()
 db.close();out(results=[dict(x) for x in rows],count=len(rows))
def op_library_stats(p):
 if not index_db().exists():out(documents=0,pdfs=0,words=0,ocrDocuments=0);return
 db=sqlite3.connect(index_db());r=db.execute("SELECT count(*),sum(CASE WHEN ext='.pdf' THEN 1 ELSE 0 END),coalesce(sum(words),0),coalesce(sum(ocr),0) FROM documents").fetchone();db.close();out(documents=r[0],pdfs=r[1]or 0,words=r[2],ocrDocuments=r[3],duckdb=duckdb is not None)
def op_capabilities(p):out(capabilities={'pdfium':True,'pdfEditing':True,'docx':Document is not None,'odt':OpenDocumentText is not None,'rtf':True,'txt':True,'legacyDocConverter':bool(find_converter()),'ocr':tess_ready(),'sqliteFts':True,'duckdb':duckdb is not None})
OPS={'info':op_info,'merge':op_merge,'extract':op_extract,'split':op_split,'delete':op_delete,'rotate':op_rotate,'reorder':op_reorder,'duplicate':op_duplicate,'insert_blank':op_insert_blank,'crop':op_crop,'watermark':op_watermark,'page_numbers':op_page_numbers,'encrypt':op_encrypt,'decrypt':op_decrypt,'text':op_text,'metadata':op_metadata,'remove_metadata':op_remove_metadata,'images_to_pdf':op_images_to_pdf,'pdf_to_images':op_pdf_to_images,'compress':op_compress,'grayscale':op_grayscale,'flatten':op_flatten,'apply_edits':op_apply_edits,'form_fields':op_form_fields,'fill_form':op_fill_form,'search_pdf':op_search_pdf,'ocr_pdf':op_ocr_pdf,'doc_open':op_doc_open,'doc_save':op_doc_save,'index_paths':op_index_paths,'search_index':op_search_index,'library_stats':op_library_stats,'capabilities':op_capabilities}
if __name__=='__main__':
 try:
  if len(sys.argv)<2:raise ValueError('Missing operation')
  op=sys.argv[1]
  if op not in OPS:raise ValueError(f'Unknown operation: {op}')
  OPS[op](payload())
 except Exception as e:fail(e)
